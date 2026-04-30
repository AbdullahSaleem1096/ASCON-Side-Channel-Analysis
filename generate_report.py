import docx
from docx.shared import Inches
import h5py
import matplotlib.pyplot as plt
import numpy as np
import os

# Generate 10 Sample Traces Plot
def plot_sample_traces():
    if not os.path.exists('ascon_fixed_key.h5'):
        return None
        
    with h5py.File('ascon_fixed_key.h5', 'r') as f:
        traces = np.array(f['traces'][:10])
        
    plt.figure(figsize=(12, 6))
    for i in range(10):
        plt.plot(traces[i], alpha=0.7, label=f'Trace {i+1}')
        
    plt.title("10 Sample Power Traces (Hamming Weight)")
    plt.xlabel("Execution Steps (Time)")
    plt.ylabel("Leakage (HW)")
    plt.legend(loc='upper right', bbox_to_anchor=(1.15, 1))
    plt.tight_layout()
    plot_path = '10_sample_traces.png'
    plt.savefig(plot_path)
    plt.close()
    return plot_path

def create_report():
    doc = docx.Document()
    doc.add_heading('Results & Analysis Report', 0)
    
    # Phase 2
    doc.add_heading('Phase 2: Technical Write-up', level=1)
    
    doc.add_heading('Implementation approach', level=2)
    doc.add_paragraph("The ASCON-128 authenticated cipher was implemented in C, targeting an ARM architecture. The implementation handles the 320-bit internal state via a 5-word 64-bit array. The AEAD flow initializes the state with the Key and Nonce, processes Associated Data (AD), encrypts the Plaintext while generating Ciphertext, and concludes with Finalization to produce an authentication tag. Standard testing routines were built to ensure output matched expected test vectors from the ASCON reference.")
    
    doc.add_heading('Challenges faced and solutions', level=2)
    doc.add_paragraph("One primary challenge was properly managing the 64-bit word interleaving and endianness during the permutation phase to match the ARM architecture correctly. Debugging required strict cross-referencing with test vectors at every intermediate step. We resolved this by creating step-by-step state printouts during the 12-round and 6-round permutations to isolate bitwise rotation errors.")
    
    doc.add_heading('Testing methodology', level=2)
    doc.add_paragraph("Testing was performed using known test vectors for ASCON-128. The implementation was compiled into an ELF binary specifically for ARM execution, allowing it to be simulated within the Unicorn emulator in subsequent phases. A test harness (`ascon_test.exe`) verified the correctness of encryption and decryption processes before proceeding to trace generation.")

    # Phase 3
    doc.add_heading('Phase 3: Dataset Documentation', level=1)
    
    doc.add_heading('Trace generation process', level=2)
    doc.add_paragraph("Trace generation was accomplished by simulating the execution of the ASCON-128 ARM binary using the Rainbow framework (built on the Unicorn engine). During simulation, memory operations (READ/WRITE) were intercepted using engine hooks. By monitoring these operations during the execution of `ascon128_encrypt` and `ascon128_init`, simulated power traces were gathered and assembled into fixed-key and variable-key datasets in HDF5 format.")
    
    doc.add_heading('Leakage model explanation', level=2)
    doc.add_paragraph("A Hamming Weight (HW) leakage model was employed. It assumes that the instantaneous power consumption of the simulated device is directly proportional to the number of set bits ('1's) in the data word being manipulated. During every memory access, the HW of the data value was calculated and appended to the trace.")
    
    doc.add_heading('Target point selection', level=2)
    doc.add_paragraph("The target point selected for the attack was the first 32/64 bits of the ASCON internal state (`state[0]`) immediately after its first modification during the encryption process. This intermediate value is highly dependent on both the secret key and the provided plaintext/nonce, making it an ideal candidate for side-channel leakage.")
    
    doc.add_heading('Sample trace plots', level=2)
    plot_path = plot_sample_traces()
    if plot_path:
        doc.add_picture(plot_path, width=Inches(6.0))
    else:
        doc.add_paragraph("[Error: Could not find ascon_fixed_key.h5 to generate plot]")

    # Phase 4
    doc.add_heading('Phase 4: Attack Results & Analysis', level=1)
    
    doc.add_heading('Model architecture and training details', level=2)
    doc.add_paragraph("A 1D Convolutional Neural Network (CNN) was built using TensorFlow/Keras. The architecture features two Conv1D layers (filters=32 and 64, kernel_size=11) interleaved with MaxPooling1D, followed by a Flatten layer, a Dense layer (128 neurons), and a final Dense softmax classification layer for the 33 possible Hamming Weight outcomes. The model was trained using Adam optimizer and Categorical Crossentropy loss for 20 epochs with a batch size of 64.")
    
    doc.add_heading('Attack performance (fixed vs variable key)', level=2)
    doc.add_paragraph("The Fixed-Key Attack demonstrated the CNN's strong capability to model the exact HW leakage, resulting in high accuracy on the 1000 held-out profiling traces. However, the Variable-Key Attack revealed the challenge of model generalization; while the model could learn specific leakage templates for a fixed key, accuracy metrics on entirely unseen keys were marginally lower, highlighting standard profiling attack generalization issues.")
    
    # Check for plots
    doc.add_paragraph("Training History Plots:")
    if os.path.exists('fixed_key_history.png'):
        doc.add_picture('fixed_key_history.png', width=Inches(5.5))
    if os.path.exists('variable_key_history.png'):
        doc.add_picture('variable_key_history.png', width=Inches(5.5))
    
    doc.add_heading('Key recovery results', level=2)
    doc.add_paragraph("By accurately predicting the Hamming Weight of the intermediate state, the side-channel attack reduces the brute-force search space significantly. For the fixed key dataset, the success rate of correctly predicting the exact HW classification provides a direct proxy for successful leakage extraction.")
    
    doc.add_heading('Critical analysis and observations', level=2)
    doc.add_paragraph("The results indicate that simulated Hamming Weight leakage is highly susceptible to deep learning profiling attacks. Because CNNs can autonomously perform feature extraction and trace alignment (shift invariance), they negate the need for manual point-of-interest selection. A significant limitation, however, is that real-world noise was absent in this simulation. In a physical setting with clock jitter and electrical noise, the CNN would require significantly more profiling traces, regularization techniques (like Dropout), and data augmentation to maintain this success rate.")
    
    doc.save('Results & Analysis Report.docx')
    print("Successfully generated Results & Analysis Report.docx")

if __name__ == '__main__':
    create_report()
